#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Criar documento
doc = Document()

# Configurar título
title = doc.add_heading('Concurso Público IASES 2022', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('AGENTE SOCIOEDUCATIVO - MASCULINO', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adicionar informações da prova
info = doc.add_paragraph()
info.add_run('Conhecimentos Específicos - Questões 57 a 70 (Questões Finais)\n').bold = True
info.add_run('Simulado de Preparação').italic = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('_' * 80)

# ============= CONHECIMENTOS ESPECÍFICOS =============
doc.add_heading('Conhecimentos Específicos', level=2)

# ============= QUESTÃO 57 =============
doc.add_heading('Questão 57', level=3)
q57 = doc.add_paragraph('Dentre as opções citadas abaixo, assinale a alternativa que corresponde a Lei, que institui o Sistema Nacional de Atendimento Socioeducativo (Sinase) e regulamenta a execução das medidas destinadas a adolescente que praticam ato infracional.\n')

q57_a = doc.add_paragraph('(A) Lei Nº 12.594, de 18 de janeiro de 2012.', style='List Number')
q57_b = doc.add_paragraph('(B) Lei Nº 8.080, de 28 de setembro de 2000.')
q57_c = doc.add_paragraph('(C) Lei Nº 11.294, de 18 de janeiro de 2002.')
q57_d = doc.add_paragraph('(D) Lei Nº 12.459, de 28 de janeiro de 2012.')
q57_e = doc.add_paragraph('(E) Lei Nº 15.294, de 18 de janeiro de 2002.')

doc.add_paragraph('')

# ============= QUESTÃO 58 =============
doc.add_heading('Questão 58', level=3)
q58 = doc.add_paragraph('Referente ao direito à profissionalização e à proteção no trabalho, previsto no Estatuto da Criança e do Adolescente, analise as afirmativas abaixo.\n\n')

q58.add_run('I. A proteção ao trabalho dos adolescentes é regulada por legislação especial.\n\n')
q58.add_run('II. Ao adolescente aprendiz, maior de quatorze anos, são assegurados os direitos trabalhistas e previdenciários.\n\n')
q58.add_run('III. Ao adolescente empregado, aprendiz, em regime familiar de trabalho, aluno de escola técnica, assistido em entidade governamental ou não-governamental, é permitido o trabalho noturno até as vinte e três horas.\n\n')
q58.add_run('É CORRETO o que se afirma em:\n')

q58_a = doc.add_paragraph('(A) I, II e III.', style='List Number')
q58_b = doc.add_paragraph('(B) I e II, apenas.')
q58_c = doc.add_paragraph('(C) III, apenas.')
q58_d = doc.add_paragraph('(D) II e III, apenas.')
q58_e = doc.add_paragraph('(E) I e III, apenas.')

doc.add_paragraph('')

# ============= QUESTÃO 59 =============
doc.add_heading('Questão 59', level=3)
q59 = doc.add_paragraph('No que diz respeito ao direito à educação, à cultura, ao esporte e ao lazer, de crianças e adolescentes, previstos no Estatuto da Criança e do Adolescente, atribua V para verdadeiro e F para falso nas afirmativas abaixo.\n\n')

q59.add_run('(__) É direito dos pais ou responsáveis ter ciência do processo pedagógico, bem como participar da definição das propostas educacionais.\n\n')
q59.add_run('(__) É dever do Estado assegurar à criança e ao adolescente, Ensino Fundamental, obrigatório e gratuito, inclusive para os que a ele não tiveram acesso na idade própria.\n\n')
q59.add_run('(__) Os pais ou responsável têm a obrigação de matricular seus filhos ou pupilos na rede regular de ensino.\n\n')
q59.add_run('Assinale a alternativa que apresenta a sequência CORRETA.\n')

q59_a = doc.add_paragraph('(A) V, F, F.', style='List Number')
q59_b = doc.add_paragraph('(B) V, V, V.')
q59_c = doc.add_paragraph('(C) F, F, F.')
q59_d = doc.add_paragraph('(D) F, V, F.')
q59_e = doc.add_paragraph('(E) V, F, V.')

doc.add_paragraph('')

# ============= QUESTÃO 60 =============
doc.add_heading('Questão 60', level=3)
q60 = doc.add_paragraph('Conforme previsto na Lei complementar Nº 706 de 2013, são atribuições dos Agentes Socioeducativos, EXCETO:\n')

q60_a = doc.add_paragraph('(A) Vistoriar periodicamente os alojamentos.', style='List Number')
q60_b = doc.add_paragraph('(B) Coletar material e/ou acompanhar o Socioeducando para exames laboratoriais.')
q60_c = doc.add_paragraph('(C) Despertar (acordar) os Socioeducandos.')
q60_d = doc.add_paragraph('(D) Procurar sempre atualizar-se em assuntos referentes à educação de Socioeducandos.')
q60_e = doc.add_paragraph('(E) Participar com os Socioeducandos, das atividades de esporte, cultura e lazer.')

doc.add_paragraph('')

# ============= QUESTÃO 61 =============
doc.add_heading('Questão 61', level=3)
q61 = doc.add_paragraph('A Declaração Universal dos Direitos Humanos (DUDH) é um documento marco na história dos direitos humanos. Elaborada por representantes de diferentes origens jurídicas e culturais de todas as regiões do mundo.\n\n')
q61.add_run('Fonte: brasil.un.org\n\n').italic = True
q61.add_run('De acordo com a Declaração Universal dos Direitos Humanos, é INCORRETO afirmar que:\n')

q61_a = doc.add_paragraph('(A) Toda pessoa acusada de um ato delituoso presume-se inocente até que a sua culpabilidade fique legalmente provada no decurso de um processo público em que todas as garantias necessárias de defesa lhe sejam asseguradas.', style='List Number')
q61_b = doc.add_paragraph('(B) Todos os indivíduos têm direito ao reconhecimento, em todos os lugares, da sua personalidade jurídica.')
q61_c = doc.add_paragraph('(C) Toda pessoa tem direito, em plena igualdade, a que a sua causa seja equitativa e publicamente julgada por um tribunal dependente e parcial que decida dos seus direitos e obrigações ou das razões de qualquer acusação em matéria penal que contra ela seja deduzida.')
q61_d = doc.add_paragraph('(D) Ninguém pode ser arbitrariamente preso, detido ou exilado.')
q61_e = doc.add_paragraph('(E) Toda pessoa sujeita a perseguição tem o direito de procurar e de beneficiar de asilo em outros países.')

doc.add_paragraph('')

# ============= QUESTÃO 62 =============
doc.add_heading('Questão 62', level=3)
q62 = doc.add_paragraph('Segundo a Declaração Universal dos Direitos Humanos (DUDH), à instrução é um direito de todo ser humano. À vista disso, atribua V para verdadeiro e F para falso nas afirmativas abaixo.\n\n')

q62.add_run('(__) A instrução deve ser gratuita, pelo menos nos graus elementares e fundamentais.\n\n')
q62.add_run('(__) A instrução elementar não deve ser de caráter obrigatório.\n\n')
q62.add_run('(__) A instrução técnico-profissional será acessível a todos, bem como a instrução superior, esta baseada no mérito.\n\n')
q62.add_run('Assinale a alternativa que apresenta a sequência CORRETA.\n')

q62_a = doc.add_paragraph('(A) V, V, V.', style='List Number')
q62_b = doc.add_paragraph('(B) F, V, F.')
q62_c = doc.add_paragraph('(C) F, V, V.')
q62_d = doc.add_paragraph('(D) V, F, V.')
q62_e = doc.add_paragraph('(E) V, F, F.')

doc.add_paragraph('')

# ============= QUESTÃO 63 =============
doc.add_heading('Questão 63', level=3)
q63 = doc.add_paragraph('Em relação ao direito da criança, do adolescente e do Jovem, previstos na Constituição Federal, atribua V para verdadeiro e F para falso nas afirmativas abaixo.\n\n')

q63.add_run('(__) São penalmente inimputáveis os menores de dezoito anos, sujeitos às normas da legislação especial.\n\n')
q63.add_run('(__) Os pais têm o dever de assistir, criar e educar os filhos menores, e os filhos maiores têm o dever de ajudar e amparar os pais na velhice, carência ou enfermidade.\n\n')
q63.add_run('(__) Os filhos, havidos ou não da relação do casamento, ou por adoção, terão os mesmos direitos e qualificações, proibidas quaisquer designações discriminatórias relativas à filiação.\n\n')
q63.add_run('Assinale a alternativa que apresenta a sequência CORRETA.\n')

q63_a = doc.add_paragraph('(A) F, V, F.', style='List Number')
q63_b = doc.add_paragraph('(B) V, V, F.')
q63_c = doc.add_paragraph('(C) F, F, V.')
q63_d = doc.add_paragraph('(D) V, V, V.')
q63_e = doc.add_paragraph('(E) V, F, V.')

doc.add_paragraph('')

# ============= QUESTÃO 64 =============
doc.add_heading('Questão 64', level=3)
q64 = doc.add_paragraph('Referente aos direitos individuais dos adolescentes, previstos no Estatuto da Criança e do Adolescente, assinale a alternativa INCORRETA.\n')

q64_a = doc.add_paragraph('(A) A internação, antes da sentença do ato infracional, pode ser determinada pelo prazo máximo de noventa dias.', style='List Number')
q64_b = doc.add_paragraph('(B) O adolescente tem direito à identificação dos responsáveis pela sua apreensão, devendo ser informado acerca de seus direitos.')
q64_c = doc.add_paragraph('(C) Nenhum adolescente deve ser privado de sua liberdade senão em flagrante de ato infracional ou por ordem escrita e fundamentada da autoridade judiciária competente.')
q64_d = doc.add_paragraph('(D) A apreensão de qualquer adolescente e o local onde se encontra recolhido serão incontinenti comunicados à autoridade judiciária competente e à família do apreendido ou à pessoa por ele indicada.')
q64_e = doc.add_paragraph('(E) O adolescente civilmente identificado não será submetido a identificação compulsória pelos órgãos policiais, de proteção e judiciais, salvo para efeito de confrontação, havendo dúvida fundada.')

doc.add_paragraph('')

# ============= QUESTÃO 65 =============
doc.add_heading('Questão 65', level=3)
q65 = doc.add_paragraph('Referente as Regras Mínimas das Nações Unidas para o Tratamento de Reclusos, analise as afirmativas abaixo, no que diz respeito ao alojamento.\n\n')

q65.add_run('I. Todos os locais destinados aos reclusos, especialmente os dormitórios, devem satisfazer todas as exigências de higiene e saúde, tomando-se devidamente em consideração as condições climáticas e, especialmente, a cubicagem de ar disponível, o espaço mínimo, a iluminação, o aquecimento e a ventilação.\n\n')
q65.add_run('II. As instalações sanitárias devem ser adequadas, de maneira a que os reclusos possam efetuar as suas necessidades quando precisarem, de modo limpo e decente.\n\n')
q65.add_run('III. As celas ou locais destinados ao descanso noturno não devem ser ocupados por mais de um recluso. Se, por razões especiais, tais como excesso temporário de população prisional, for necessário que a administração prisional central adote exceções a esta regra deve evitar-se que dois reclusos sejam alojados numa mesma cela ou local.\n\n')
q65.add_run('É CORRETO o que se afirma em:\n')

q65_a = doc.add_paragraph('(A) I, II e III.', style='List Number')
q65_b = doc.add_paragraph('(B) I e III, apenas.')
q65_c = doc.add_paragraph('(C) I e II, apenas.')
q65_d = doc.add_paragraph('(D) II e III, apenas.')
q65_e = doc.add_paragraph('(E) II, apenas.')

doc.add_paragraph('')

# ============= QUESTÃO 66 =============
doc.add_heading('Questão 66', level=3)
q66 = doc.add_paragraph('Referente as visitas a adolescente em cumprimento de medida de internação, previstas na Lei do Sistema Nacional de Atendimento Socioeducativo - SINASE, atribua V para verdadeiro e F para falso nas afirmativas abaixo.\n\n')

q66.add_run('(__) Não é assegurado ao adolescente que viva, comprovadamente, em união estável o direito à visita íntima.\n\n')
q66.add_run('(__) É garantido aos adolescentes em cumprimento de medida socioeducativa de internação o direito de receber visita dos filhos, independentemente da idade desses.\n\n')
q66.add_run('(__) O regulamento interno estabelecerá as hipóteses de proibição da entrada de objetos na unidade de internação, vedando o acesso aos seus portadores.\n\n')
q66.add_run('Assinale a alternativa que apresenta a sequência CORRETA.\n')

q66_a = doc.add_paragraph('(A) F, F, V.', style='List Number')
q66_b = doc.add_paragraph('(B) F, F, F.')
q66_c = doc.add_paragraph('(C) F, V, V.')
q66_d = doc.add_paragraph('(D) V, F, V.')
q66_e = doc.add_paragraph('(E) V, V, V.')

doc.add_paragraph('')

# ============= QUESTÃO 67 =============
doc.add_heading('Questão 67', level=3)
obs67 = doc.add_paragraph()
obs67.add_run('(Questão anulada)\n\n').italic = True
q67 = doc.add_paragraph('Em relação as atitudes e posturas que influenciam no compromisso ético dos profissionais no Serviço Público, atribua V para verdadeiro e F para falso nas afirmativas abaixo.\n\n')

q67.add_run('(__) O decoro é uma "postura" porque une a disposição interna para agir corretamente com a aparência desse agir. Decoro, do latim decorum, é "a face pública de um estado pessoal da honradez" (David Burchell).\n\n')
q67.add_run('(__) A objetividade significa uma abordagem razoavelmente distanciada e serena do trabalho a fazer. Isso não significa indiferença ou frieza: trata-se apenas de evitar que sentimentos explosivos atrapalhem o nosso desempenho (ENAP, 2014).\n\n')
q67.add_run('(__) A civilidade significa disposição para justificar publicamente decisões tomadas ou estratégias adotadas, e abertura para ouvir interpelações, críticas e sugestões. Porém, de forma respeitosa, independentemente da simpatia pessoal que se tenha pelo interlocutor desempenho (ENAP, 2014).\n\n')
q67.add_run('Fonte: repositorio.enap.gov.br\n\n').italic = True
q67.add_run('Assinale a alternativa que apresenta a sequência CORRETA.\n')

q67_a = doc.add_paragraph('(A) F, V, F.', style='List Number')
q67_b = doc.add_paragraph('(B) F, F, F.')
q67_c = doc.add_paragraph('(C) F, V, V.')
q67_d = doc.add_paragraph('(D) V, F, F.')
q67_e = doc.add_paragraph('(E) V, V, V.')

doc.add_paragraph('')

# ============= QUESTÃO 68 =============
doc.add_heading('Questão 68', level=3)
q68 = doc.add_paragraph('Referente as Regras Mínimas das Nações Unidas, quanto ao serviço médico para o Tratamento de Reclusos, atribua V para verdadeiro e F para falso nas afirmativas abaixo.\n\n')

q68.add_run('(__) Os serviços de saúde devem elaborar registos médicos individuais, confidenciais, atualizados e precisos para cada um dos reclusos, que a eles devem ter acesso, sempre que solicitado. O recluso não pode ter acesso ao seu registo médico, mesmo através de uma terceira pessoa por si designada.\n\n')
q68.add_run('(__) As decisões clínicas só podem ser tomadas por profissionais de saúde responsáveis e não podem ser modificadas ou ignoradas pela equipa prisional não médica.\n\n')
q68.add_run('(__) O registro médico deve ser encaminhado para o serviço de saúde do estabelecimento prisional para o qual o recluso é transferido, encontrando-se sujeito à confidencialidade médica.\n\n')
q68.add_run('Assinale a alternativa que apresenta a sequência CORRETA.\n')

q68_a = doc.add_paragraph('(A) F, F, V.', style='List Number')
q68_b = doc.add_paragraph('(B) F, F, F.')
q68_c = doc.add_paragraph('(C) V, F, F.')
q68_d = doc.add_paragraph('(D) V, V, V.')
q68_e = doc.add_paragraph('(E) F, V, V.')

doc.add_paragraph('')

# ============= QUESTÃO 69 =============
doc.add_heading('Questão 69', level=3)
q69 = doc.add_paragraph('No que se refere ao processo de adoção de crianças e adolescentes, previsto no Estatuto da Criança e do Adolescente, analise as afirmativas abaixo.\n\n')

q69.add_run('I. O adotando deve contar com, no máximo, vinte e um anos à data do pedido, salvo se já estiver sob a guarda ou tutela dos adotantes.\n\n')
q69.add_run('II. A adoção atribui a condição de filho ao adotado, com os mesmos direitos e deveres, inclusive sucessórios, desligando-o de qualquer vínculo com pais e parentes, salvo os impedimentos matrimoniais.\n\n')
q69.add_run('III. A adoção não depende do consentimento dos pais ou do representante legal do adotando.\n\n')
q69.add_run('É CORRETO o que se afirma em:\n')

q69_a = doc.add_paragraph('(A) II, apenas.', style='List Number')
q69_b = doc.add_paragraph('(B) III, apenas.')
q69_c = doc.add_paragraph('(C) I e III, apenas.')
q69_d = doc.add_paragraph('(D) I, II e III.')
q69_e = doc.add_paragraph('(E) I e II, apenas.')

doc.add_paragraph('')

# ============= QUESTÃO 70 =============
doc.add_heading('Questão 70', level=3)
q70 = doc.add_paragraph('No que se refere ao Funcionamento do Sistema Único de Segurança Pública (SUSP), é CORRETO afirmar que:\n')

q70_a = doc.add_paragraph('(A) A União poderá apoiar os Estados, o Distrito Federal e os Municípios, quando não dispuserem de condições técnicas e operacionais necessárias à implementação do SUSP.', style='List Number')
q70_b = doc.add_paragraph('(B) O SUSP será coordenado pelo Ministério de Desenvolvimento e Assistência Social.')
q70_c = doc.add_paragraph('(C) As aquisições de bens e serviços para os órgãos integrantes do SUSP terão por objetivo a eficácia de suas atividades, não sendo necessário obedecer a critérios técnicos de modernidade, eficiência e resistência, observadas as normas de licitação e contratos.')
q70_d = doc.add_paragraph('(D) Os órgãos integrantes do SUSP não devem atuar em rodovias, terminais rodoviários, ferrovias e hidrovias federais, portos e aeroportos, no âmbito das respectivas competências, em efetiva integração com o órgão cujo local de atuação, mesmo que, esteja sob sua circunscrição.')
q70_e = doc.add_paragraph('(E) O compartilhamento de informações não deve ser feito por meio eletrônico.')

# Salvar documento
doc.save('/home/user/iases/Prova_Simulado_Questoes_57-70.docx')
print("Documento criado com sucesso: Prova_Simulado_Questoes_57-70.docx")
