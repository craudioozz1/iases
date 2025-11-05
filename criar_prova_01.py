#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Criar documento
doc = Document()

# Configurar título
title = doc.add_heading('Concurso Público IASES 2022', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('AGENTE SOCIOEDUCATIVO - MASCULINO', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adicionar informações da prova
info = doc.add_paragraph()
info.add_run('Língua Portuguesa - Questões 1 a 14\n').bold = True
info.add_run('Simulado de Preparação').italic = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('_' * 80)

# ============= TEXTO BASE QUESTÕES 1 A 5 =============
doc.add_heading('Texto Base - Questões 1 a 5', level=2)

texto1 = doc.add_paragraph()
texto1.add_run('Declaração Universal dos Direitos Humanos\n\n').bold = True
texto1.add_run('Adotada e proclamada pela Assembleia Geral das Nações Unidas (resolução 217 A III) em 10 de dezembro 1948.\n\n')
texto1.add_run('Preâmbulo.\n\n').bold = True
texto1.add_run('Considerando que o reconhecimento da dignidade inerente a todos os membros da família humana e de seus direitos iguais e inalienáveis é o fundamento da liberdade, da justiça e da paz no mundo; [...] A Assembleia Geral proclama a presente Declaração Universal dos Direitos Humanos como o ideal comum a ser atingido por todos os povos e todas as nações, com o objetivo de que cada indivíduo e cada órgão da sociedade tendo sempre em mente esta Declaração, esforce-se, por meio do ensino e da educação, por promover o respeito a esses direitos e liberdades, e, pela adoção de medidas progressivas de caráter nacional e internacional, por assegurar o seu reconhecimento e a sua observância universais e efetivos, tanto entre os povos dos próprios Países-Membros quanto entre os povos dos territórios sob sua jurisdição. [...]\n\n')
texto1.add_run('Disponível em: https://brasil.un.org/ Acesso em: 25/01/2023').italic = True

doc.add_paragraph('')

# ============= QUESTÃO 01 =============
doc.add_heading('Questão 01', level=3)
q1 = doc.add_paragraph('Em relação às regras de acentuação gráfica, o que justifica a acentuação das palavras INALIENÁVEIS, INDIVÍDUO e OBSERVÂNCIA é:\n')

q1_a = doc.add_paragraph('(A) São acentuadas todas as palavras paroxítonas terminadas em -um e -uns.', style='List Number')
q1_b = doc.add_paragraph('(B) Acentuam-se as palavras paroxítonas terminadas em ditongos.')
q1_c = doc.add_paragraph('(C) São acentuadas as palavras paroxítonas terminadas em -us.')
q1_d = doc.add_paragraph('(D) Acentuam-se as palavras paroxítonas terminadas em -i seguido ou não de -s, são graficamente acentuadas.')
q1_e = doc.add_paragraph('(E) Recebem acento gráfico todas as paroxítonas que têm terminação -om ou -ons.')

doc.add_paragraph('')

# ============= QUESTÃO 02 =============
doc.add_heading('Questão 02', level=3)
q2 = doc.add_paragraph('A função da linguagem presente no texto é:\n')

q2_a = doc.add_paragraph('(A) Conativa ou apelativa.', style='List Number')
q2_b = doc.add_paragraph('(B) Emotiva ou expressiva.')
q2_c = doc.add_paragraph('(C) Poética.')
q2_d = doc.add_paragraph('(D) Referencial ou denotativa.')
q2_e = doc.add_paragraph('(E) Metalinguística.')

doc.add_paragraph('')

# ============= QUESTÃO 03 =============
doc.add_heading('Questão 03', level=3)
q3 = doc.add_paragraph('A palavra ASSEMBLEIA é grafada sem acento porque:\n')

q3_a = doc.add_paragraph('(A) As formas verbais que possuem o acento tônico na raiz, com o (u) tônico precedido de "g" ou "q" e seguido de "e" ou "i" não serão mais acentuadas.', style='List Number')
q3_b = doc.add_paragraph('(B) Não serão acentuadas as palavras em que o "i" e o "u" tônicos dos hiatos, são antecedidos de ditongos.')
q3_c = doc.add_paragraph('(C) As palavras com ditongos terminados em -ei e -oi, não são mais acentuados.')
q3_d = doc.add_paragraph('(D) O acento agudo não será mais usado para diferenciar determinados vocábulos.')
q3_e = doc.add_paragraph('(E) Não serão mais acentuadas as vogais tônicas quando formarem hiato.')

doc.add_paragraph('')

# ============= QUESTÃO 04 =============
doc.add_heading('Questão 04', level=3)
q4 = doc.add_paragraph('O uso das reticências no excerto do preâmbulo da Declaração Universal dos Direitos Humanos expressa:\n')

q4_a = doc.add_paragraph('(A) Isolar palavras, frases intercaladas de caráter explicativo e datas.', style='List Number')
q4_b = doc.add_paragraph('(B) Indicar a presença de apostos ou orações apositivas, enumerações ou sequência de palavras que explicam, resumem ideias anteriores.')
q4_c = doc.add_paragraph('(C) Separar períodos não relacionados entre si.')
q4_d = doc.add_paragraph('(D) Indicar supressão de palavra (s) numa frase transcrita.')
q4_e = doc.add_paragraph('(E) Intenção de sugerir prolongamento de ideia, omissão de trechos do texto original, indicar dúvidas ou hesitação do falante.')

doc.add_paragraph('')

# ============= QUESTÃO 05 =============
doc.add_heading('Questão 05', level=3)
q5 = doc.add_paragraph('O que se pode inferir por "direitos inalienáveis" citado no texto?\n')

q5_a = doc.add_paragraph('(A) Os direitos inalienáveis são as manifestações de direitos de cidadãos de algumas partes do mundo.', style='List Number')
q5_b = doc.add_paragraph('(B) Os direitos inalienáveis são aqueles iguais para todos e que podem ser cedidos a outrem.')
q5_c = doc.add_paragraph('(C) Os direitos inalienáveis são todos os direitos de algumas etnias específicas.')
q5_d = doc.add_paragraph('(D) Os direitos inalienáveis são aqueles destinados a alguns cidadãos.')
q5_e = doc.add_paragraph('(E) Os direitos inalienáveis são todos os direitos fundamentais que não podem ser vendidos ou cedidos.')

doc.add_page_break()

# ============= TEXTO BASE QUESTÕES 6 A 10 =============
doc.add_heading('Texto Base - Questões 6 a 10', level=2)

texto2 = doc.add_paragraph()
texto2.add_run('REFLEXÃO SOBRE O "ESTATUTO DA CRIANÇA E DO ADOLESCENTE"\n\n').bold = True
texto2.add_run('(1º) O Estatuto da Criança e do Adolescente (ECA - Lei 8069/90) foi fruto da necessidade da criação de uma Justiça especializada e cujo objetivo é de julgar as infrações cometidas pelos adolescentes entre doze e dezoito anos (artigo 2º) do ECA.\n\n')
texto2.add_run('(2º) O dicionário de Aurélio Buarque de Holanda conceitua o vocábulo adolescente como: aquele que está no começo, no início, que ainda não atingiu todo o vigor. Adolescentes são pessoas ainda em formação, cuja estrutura física e psíquica não atingiu sua plenitude, bem como a sua personalidade. Sendo assim, são pessoas especiais que merecem a criação de uma Justiça especializada, diferenciada daquela utilizada para adultos, haja vista, suas diferenças.\n\n')
texto2.add_run('(3º) Como seres especiais, cuja personalidade, intelecto, caráter estão ainda em formação a tarefa de redirecioná-los e reeducá-los é mais branda e menos trabalhosa, pois as crianças e os adolescentes são mais suscetíveis em assimilar as ditas orientações. Pense nisso, redobre sua atenção para com os seres humanos em formação! Eles merecem o nosso carinho!\n\n')
texto2.add_run('(4º) O ECA, portanto, prevê um tratamento diferenciado para os adolescentes infratores, classificando-os como pessoas especiais de direitos, procurando garantir que sua formação seja sólida e harmoniosa perante a sociedade, garantindo assim a retomada de uma vida social plena sem problemas ou incidentes, lastreados em valores éticos, sociais e familiares, afastando-os de uma vida pregressa gregária que não deve prevalecer, em nenhuma hipótese durante ao seu desenvolvimento, sob pena de se tornar um doente incurável.\n\n')
texto2.add_run('(Reflexões sobre o Estatuto da Criança e do Adolescente - Jus.com.br | Jus Navigandi) - (Adaptado)').italic = True

doc.add_paragraph('')

# ============= QUESTÃO 06 =============
doc.add_heading('Questão 06', level=3)
q6 = doc.add_paragraph('Marque a "Função da linguagem" que predomina no trecho do texto: "Pense nisso, redobre sua atenção para com os seres humanos em formação, pois eles merecem o nosso carinho!"\n')

q6_a = doc.add_paragraph('(A) Metalinguística.', style='List Number')
q6_b = doc.add_paragraph('(B) Fática.')
q6_c = doc.add_paragraph('(C) Emotiva.')
q6_d = doc.add_paragraph('(D) Apelativa.')
q6_e = doc.add_paragraph('(E) Poética.')

doc.add_paragraph('')

# ============= QUESTÃO 07 =============
doc.add_heading('Questão 07', level=3)
q7 = doc.add_paragraph('Sobre os componentes linguísticos do período transcrito a seguir, marque a alternativa com análise CORRETA.\n\n')
q7.add_run('"Como seres especiais, cuja personalidade, intelecto, caráter estão ainda em formação a tarefa de redirecioná-los e reeducá-los é mais branda e menos trabalhosa, as crianças e os adolescentes são mais suscetíveis em assimilar as ditas orientações."\n').italic = True

q7_a = doc.add_paragraph('(A) A palavra dissílaba: mais faz antítese com muitos.', style='List Number')
q7_b = doc.add_paragraph('(B) No período, temos exemplos de expressões incoerentes com o tema do texto.')
q7_c = doc.add_paragraph('(C) Os verbos: redirecioná-los e reeducá-los se identificam somente porque têm a mesma quantidade de sílabas e por serem oxítonos.')
q7_d = doc.add_paragraph('(D) No período, existem vários exemplos de comparações.')
q7_e = doc.add_paragraph('(E) A organização do período enunciado comprova o uso da figura de linguagem denominada "Hipérbato".')

doc.add_paragraph('')

# ============= QUESTÃO 08 =============
doc.add_heading('Questão 08', level=3)
q8 = doc.add_paragraph('Analise as assertivas com V(Verdadeiro) ou F(Falso).\n')

q8.add_run('(__) No texto, temos exemplo de discurso dissertativo coerente e coeso com o tema apresentado.\n\n')
q8.add_run('(__) O período do (2º) "O dicionário de Aurélio Buarque de Holanda conceitua o vocábulo adolescente como: aquele que está no começo, no início, que ainda não atingiu todo o vigor" - está escrito com dois pontos, para introduzir um esclarecimento, as duas primeiras vírgulas separam expressão que tem a mesma função de "no começo".\n\n')
q8.add_run('(__) Uso de termos oxítonos, paroxítonos e proparoxítonos, conforme se pode exemplificar pela série: "atenção, caráter, hipótese".\n\n')
q8.add_run('(__) Uso de figuras de linguagem, como metáforas e comparações para conceituar o conteúdo do "ECA", conforme está escrito no (2º) do texto.\n\n')
q8.add_run('Marque a alternativa com a série CORRETA:\n')

q8_a = doc.add_paragraph('(A) F, V, F, V.', style='List Number')
q8_b = doc.add_paragraph('(B) V, V, V, V.')
q8_c = doc.add_paragraph('(C) V, V, V, F.')
q8_d = doc.add_paragraph('(D) V, V, F, F.')
q8_e = doc.add_paragraph('(E) V, V, F, F.')

doc.add_paragraph('')

# ============= QUESTÃO 09 =============
doc.add_heading('Questão 09', level=3)
q9 = doc.add_paragraph('Analise as assertivas com V(Verdadeiro) ou F(Falso).\n')

q9.add_run('(__) No texto, predomina linguagem denotativa.\n\n')
q9.add_run('(__) Verbete de dicionário é um texto que procura, da forma mais sintética e objetiva possível, definir uma palavra, dando conta de suas mais variadas acepções, este gênero textual relaciona-se com propriedade ao texto "Reflexão sobre o Estatuto da Criança e do Adolescente".\n\n')
q9.add_run('(__) Uso de discurso direto, exemplificado pela frase: "Pense nisso, redobre sua atenção para com os seres humanos em formação!"\n\n')
q9.add_run('(__) Uso de discurso indireto predominante no texto, exemplo: "O ECA prevê um tratamento diferenciado para os adolescentes infratores".\n\n')
q9.add_run('Marque a alternativa com a sequência CORRETA da sua análise:\n')

q9_a = doc.add_paragraph('(A) V, V, F, F.', style='List Number')
q9_b = doc.add_paragraph('(B) V, V, V, V.')
q9_c = doc.add_paragraph('(C) V, V, V, F.')
q9_d = doc.add_paragraph('(D) F, V, V, V.')
q9_e = doc.add_paragraph('(E) V, F, V, F.')

doc.add_paragraph('')

# ============= QUESTÃO 10 =============
doc.add_heading('Questão 10', level=3)
q10 = doc.add_paragraph('Marque o que não se comprova na estrutura textual.\n')

q10_a = doc.add_paragraph('(A) Foco temático no Estatuto da Criança e do Adolescente.', style='List Number')
q10_b = doc.add_paragraph('(B) Texto característico de dissertação em terceira pessoa.')
q10_c = doc.add_paragraph('(C) Texto escrito com predominância de termos denotativos.')
q10_d = doc.add_paragraph('(D) Uso de conceitos extraídos de um dicionário, que representa o título bibliográfico da Língua Portuguesa, publicado por um autor consagrado nacional e internacionalmente.')
q10_e = doc.add_paragraph('(E) Narração bem detalhada, apresentada por um narrador observador.')

doc.add_page_break()

# ============= QUESTÕES INDIVIDUAIS 11 A 14 =============
doc.add_heading('Questões 11 a 14', level=2)

# ============= QUESTÃO 11 =============
doc.add_heading('Questão 11', level=3)
q11 = doc.add_paragraph('Assinale a única alternativa que está de acordo com as normas de regência da língua culta:\n')

q11_a = doc.add_paragraph('(A) Notifiquei-o de que não pretendia substituí-lo na liderança do batalhão, porque apesar de ter sempre servido à polícia, jamais aspirei a tal posição.', style='List Number')
q11_b = doc.add_paragraph('(B) Notifiquei-lhe de que não pretendia substituir-lhe na liderança do batalhão, porque apesar de ter sempre servido à polícia, jamais aspirei a tal posição.')
q11_c = doc.add_paragraph('(C) Notifiquei-o de que não pretendia substituir-lhe na liderança do batalhão, porque apesar de ter sempre servido à polícia, jamais aspirei tal posição.')
q11_d = doc.add_paragraph('(D) Notifiquei-lhe de que não pretendia substituí-lo na liderança do batalhão, porque apesar de ter sempre servido a polícia, jamais aspirei a tal posição.')
q11_e = doc.add_paragraph('(E) Notifiquei-o de que não pretendia substituí-lo na liderança do batalhão, porque apesar de ter sempre servido a polícia, jamais aspirei tal posição.')

doc.add_paragraph('')

# ============= QUESTÃO 12 =============
doc.add_heading('Questão 12', level=3)
q12 = doc.add_paragraph('Assinale a alternativa em que o adjetivo que qualifica o substantivo seja explicativo:\n')

q12_a = doc.add_paragraph('(A) Cão faminto.', style='List Number')
q12_b = doc.add_paragraph('(B) Jovem estudioso.')
q12_c = doc.add_paragraph('(C) Homem mortal.')
q12_d = doc.add_paragraph('(D) Dia quente.')
q12_e = doc.add_paragraph('(E) Lua crescente.')

doc.add_paragraph('')

# ============= QUESTÃO 13 =============
doc.add_heading('Questão 13', level=3)
q13 = doc.add_paragraph('Assinale a alternativa que contém o emprego INADEQUADO de pronome pessoal:\n')

q13_a = doc.add_paragraph('(A) Está tudo certo entre eu e ti.', style='List Number')
q13_b = doc.add_paragraph('(B) Nem tudo está claro entre ele e você.')
q13_c = doc.add_paragraph('(C) Nada mais há entre mim e ti.')
q13_d = doc.add_paragraph('(D) Nada pode ser omitido entre ele e ela.')
q13_e = doc.add_paragraph('(E) Pedro gosta de dormir entre ela e ti.')

doc.add_paragraph('')

# ============= QUESTÃO 14 =============
doc.add_heading('Questão 14', level=3)
q14 = doc.add_paragraph('Assinale a alternativa que indica o grau correto do adjetivo nas seguinte orações:\n\n')
q14.add_run('I. Aquele jogador é habilidosíssimo.\n')
q14.add_run('II. Suco de maracujá é o melhor dos sucos.\n\n')

q14_a = doc.add_paragraph('(A) I. Superlativo absoluto analítico. II. Superlativo relativo de superioridade.', style='List Number')
q14_b = doc.add_paragraph('(B) I. Superlativo relativo de superioridade. II. Comparativo de superioridade.')
q14_c = doc.add_paragraph('(C) I. Comparativo de superioridade. II. Superlativo absoluto sintético.')
q14_d = doc.add_paragraph('(D) I. Superlativo absoluto sintético. II. Superlativo relativo de superioridade.')
q14_e = doc.add_paragraph('(E) I. Comparativo de igualdade. II. Comparativo de superioridade.')

# Salvar documento
doc.save('/home/user/iases/Prova_Simulado_Questoes_01-14.docx')
print("Documento criado com sucesso: Prova_Simulado_Questoes_01-14.docx")
