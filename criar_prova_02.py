#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Criar documento
doc = Document()

# Configurar título
title = doc.add_heading('Concurso Público IASES 2022', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('AGENTE SOCIOEDUCATIVO - MASCULINO', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adicionar informações da prova
info = doc.add_paragraph()
info.add_run('Língua Portuguesa (Q15) e Matemática (Q16-28) - Questões 15 a 28\n').bold = True
info.add_run('Simulado de Preparação').italic = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('_' * 80)

# ============= QUESTÃO 15 =============
doc.add_heading('Língua Portuguesa', level=2)
doc.add_heading('Questão 15', level=3)
q15 = doc.add_paragraph('Analise as assertivas com V(Verdadeiro) ou F(Falso).\n')

q15.add_run('(__) Na frase nominal exclamativa: "Bom-dia, jovens estudiosos!" - temos um vocativo (chamamento) representado por "jovens estudiosos".\n\n')
q15.add_run('(__) A frase: "Os cavalheiros são gentis até mesmo como cavaleiros" - está estruturada com parônimos.\n\n')
q15.add_run('(__) Na frase: "Eu almoço satisfeito, porque este almoço está muito bem feito" - temos homônimos com a mesma grafia e com pronúncias diferentes - são homônimos homógrafos heterófonos.\n\n')
q15.add_run('(__) Na frase: "Leve os garotos agora, porque o trânsito está leve" - temos homônimos perfeitos (iguais na grafia e na pronúncia).\n\n')
q15.add_run('Marque a alternativa que apresenta a sequência CORRETA da sua análise:\n')

q15_a = doc.add_paragraph('(A) F, F, V, V.', style='List Number')
q15_b = doc.add_paragraph('(B) V, V, V, V.')
q15_c = doc.add_paragraph('(C) V, F, V, V.')
q15_d = doc.add_paragraph('(D) F, V, F, V.')
q15_e = doc.add_paragraph('(E) F, V, F, F.')

doc.add_page_break()

# ============= MATEMÁTICA =============
doc.add_heading('Matemática', level=2)

# ============= QUESTÃO 16 =============
doc.add_heading('Questão 16', level=3)
q16 = doc.add_paragraph('A diretoria de uma torcida organizada formada por 400 componentes colocou em votação dois novos modelos de camisas, o modelo "regata" e o modelo "polo". 60% dos torcedores associados escolheram o modelo regata e 32% dos torcedores associados escolheram o modelo polo. Qual o total de torcedores associados que se abstiveram do voto?\n')

q16_a = doc.add_paragraph('(A) 50 torcedores.', style='List Number')
q16_b = doc.add_paragraph('(B) 15 torcedores.')
q16_c = doc.add_paragraph('(C) 40 torcedores.')
q16_d = doc.add_paragraph('(D) 32 torcedores.')
q16_e = doc.add_paragraph('(E) 30 torcedores.')

doc.add_paragraph('')

# ============= QUESTÃO 17 =============
doc.add_heading('Questão 17', level=3)
q17 = doc.add_paragraph('Luiza comprou um terreno de 384 m², com medidas dadas na imagem abaixo.\n\n')
q17.add_run('Quais são as dimensões do terreno?\n')

q17_a = doc.add_paragraph('(A) O terreno tem dimensões de 21m x 13m.', style='List Number')
q17_b = doc.add_paragraph('(B) O terreno tem dimensões de 29m x 21m.')
q17_c = doc.add_paragraph('(C) O terreno tem dimensões de 22m x 14m.')
q17_d = doc.add_paragraph('(D) O terreno tem dimensões de 30m x 22m.')
q17_e = doc.add_paragraph('(E) O terreno tem dimensões de 24m x 16m.')

doc.add_paragraph('')

# ============= QUESTÃO 18 =============
doc.add_heading('Questão 18', level=3)
obs = doc.add_paragraph()
obs.add_run('(Questão anulada)\n\n').italic = True
q18 = doc.add_paragraph('O tempo em minutos para que um anestesista administre o medicamento e leve o paciente a total narcose é dado pela função f(i) = 2 + log(i/6), onde i é a idade do paciente e f(i) é o tempo dado em minutos. Em um paciente de 30 anos, o tempo necessário para total narcose é de? (Considere log 2 = 0,3.)\n')

q18_a = doc.add_paragraph('(A) 2 minutos e 7 segundos.', style='List Number')
q18_b = doc.add_paragraph('(B) 1 minutos e 42 segundos.')
q18_c = doc.add_paragraph('(C) 2 minutos e 56 segundos.')
q18_d = doc.add_paragraph('(D) 2 minutos e 42 segundos.')
q18_e = doc.add_paragraph('(E) 1 minuto e 42 segundos.')

doc.add_paragraph('')

# ============= QUESTÃO 19 =============
doc.add_heading('Questão 19', level=3)
q19 = doc.add_paragraph('A área de um retângulo é igual a 28 cm², sendo a base igual x - 1 cm e a altura igual x - 4 cm, onde x > 0. Assinale a alternativa que apresenta, em centímetros, a base e a altura deste retângulo.\n')

q19_a = doc.add_paragraph('(A) base = 16 cm e altura = 12 cm', style='List Number')
q19_b = doc.add_paragraph('(B) base = 7 cm e altura = 4 cm')
q19_c = doc.add_paragraph('(C) base = 14 cm e altura = 2 cm')
q19_d = doc.add_paragraph('(D) base = 20 cm e altura = 8 cm')
q19_e = doc.add_paragraph('(E) base = - 4 cm e altura = - 7 cm')

doc.add_paragraph('')

# ============= QUESTÃO 20 =============
doc.add_heading('Questão 20', level=3)
q20 = doc.add_paragraph('O setor de suporte da TI possui três caixas com conectores de rede do tipo RJ45, e cada caixa possui dois compartimentos. A caixa 1 contém 50 conectores RJ45 em cada compartimento. A caixa 2 contém 10 conectores RJ45 em cada compartimento e a caixa 3 contém 50 conectores RJ45 em um compartimento e 10 conectores RJ45 no outro compartimento. Escolhendo uma caixa ao acaso e abrindo um compartimento, se for encontrado 50 conectores RJ45, qual é a probabilidade de que no outro compartimento seja encontrado, também, 50 conectores RJ45?\n')

q20_a = doc.add_paragraph('(A) 3/4 ou 75%.', style='List Number')
q20_b = doc.add_paragraph('(B) 3/5 ou 60%.')
q20_c = doc.add_paragraph('(C) 2/3 ou 66,67%.')
q20_d = doc.add_paragraph('(D) 1/3 ou 33,33%.')
q20_e = doc.add_paragraph('(E) 1/2 ou 50%.')

doc.add_paragraph('')

# ============= QUESTÃO 21 =============
doc.add_heading('Questão 21', level=3)
q21 = doc.add_paragraph('O triângulo é uma figura geométrica que ocupa o espaço interno limitado por três segmentos de reta que concorrem, dois a dois, em três pontos diferentes formando três lados e três ângulos internos que somam 180º. Dado um triângulo que possui os seguintes ângulos: (x + 2)º, (3x - 25)º e (x + 98)º. Assinale a alternativa que apresenta corretamente as medidas em graus destes ângulos:\n')

q21_a = doc.add_paragraph('(A) 27º; 38º e 115º', style='List Number')
q21_b = doc.add_paragraph('(B) 23º; 42º e 115º')
q21_c = doc.add_paragraph('(C) 28º; 33º e 119º')
q21_d = doc.add_paragraph('(D) 21º; 40º e 119º')
q21_e = doc.add_paragraph('(E) 23º; 38º e 119º')

doc.add_paragraph('')

# ============= QUESTÃO 22 =============
doc.add_heading('Questão 22', level=3)
q22 = doc.add_paragraph('Foi instalado no paciente um soro de 600 ml com o gotejamento de 5 (cinco gotas) por segundo. Considerando que uma gota de soro é formada em média de 5 x 10⁻² ml, quanto tempo, em minutos, será necessário para todo o soro ser totalmente administrado?\n')

q22_a = doc.add_paragraph('(A) 50 minutos.', style='List Number')
q22_b = doc.add_paragraph('(B) 30 minutos.')
q22_c = doc.add_paragraph('(C) 40 minutos.')
q22_d = doc.add_paragraph('(D) 60 minutos.')
q22_e = doc.add_paragraph('(E) 20 minutos.')

doc.add_paragraph('')

# ============= QUESTÃO 23 =============
doc.add_heading('Questão 23', level=3)
q23 = doc.add_paragraph('Júlio lançou dois dados ao mesmo tempo sobre uma mesa. Qual a probabilidade de dois números iguais ficarem voltados para cima?\n')

q23_a = doc.add_paragraph('(A) A probabilidade é de aproximadamente 30,7%.', style='List Number')
q23_b = doc.add_paragraph('(B) A probabilidade é de aproximadamente 23,45%.')
q23_c = doc.add_paragraph('(C) A probabilidade é de aproximadamente 12,22%.')
q23_d = doc.add_paragraph('(D) A probabilidade é de aproximadamente 29,05%.')
q23_e = doc.add_paragraph('(E) A probabilidade é de aproximadamente 16,66%.')

doc.add_paragraph('')

# ============= QUESTÃO 24 =============
doc.add_heading('Questão 24', level=3)
q24 = doc.add_paragraph('Angela, Claudia e Lúcia trabalham juntas produzindo tapetes artesanais. Em um determinado mês elas faturaram R$ 2.800,00, mas como Claudia tinha trabalhado o dobro do tempo de cada uma das outras duas, ficou com uma parte dos lucros proporcional à sua dedicação. Quanto Angela recebeu?\n')

q24_a = doc.add_paragraph('(A) Angela recebeu R$ 600,00.', style='List Number')
q24_b = doc.add_paragraph('(B) Angela recebeu R$ 800,00.')
q24_c = doc.add_paragraph('(C) Angela recebeu R$ 650,00.')
q24_d = doc.add_paragraph('(D) Angela recebeu R$ 700,00.')
q24_e = doc.add_paragraph('(E) Angela recebeu R$ 550,00.')

doc.add_paragraph('')

# ============= QUESTÃO 25 =============
doc.add_heading('Questão 25', level=3)
q25 = doc.add_paragraph('Se x é um número que pertence ao conjunto dos números racionais, qual das alternativas abaixo NÃO é verdadeira?\n')

q25_a = doc.add_paragraph('(A) x = 1,238761...', style='List Number')
q25_b = doc.add_paragraph('(B) x = 9/2')
q25_c = doc.add_paragraph('(C) x = 27')
q25_d = doc.add_paragraph('(D) x = 0,333333...')
q25_e = doc.add_paragraph('(E) x = - 45')

doc.add_paragraph('')

# ============= QUESTÃO 26 =============
doc.add_heading('Questão 26', level=3)
q26 = doc.add_paragraph('Após pesquisar em várias lojas de automóveis usados, João encontrou um carro em perfeitas condições por R$ 85.000,00 a vista. Como João não possuía esta quantia ele acertou a compra do veículo em duas prestações de R$ 45.000,00, uma no ato da compra e outra um mês depois. Qual a taxa de juros mensal que a loja cobrou nessa operação?\n')

q26_a = doc.add_paragraph('(A) 10%', style='List Number')
q26_b = doc.add_paragraph('(B) 11,5%')
q26_c = doc.add_paragraph('(C) 12,5%')
q26_d = doc.add_paragraph('(D) 8,5%')
q26_e = doc.add_paragraph('(E) 12%')

doc.add_paragraph('')

# ============= QUESTÃO 27 =============
doc.add_heading('Questão 27', level=3)
q27 = doc.add_paragraph('Fernanda é a última de uma fila onde 45 pessoas, com ela, esperam ser atendidas. Se o atendimento inicia em um momento 0 e a cada 2 minutos uma pessoa é chamada, em quanto tempo Fernanda será atendida?\n')

q27_a = doc.add_paragraph('(A) Em 1h28min.', style='List Number')
q27_b = doc.add_paragraph('(B) Em 2h05min.')
q27_c = doc.add_paragraph('(C) Em 1h45min.')
q27_d = doc.add_paragraph('(D) Em 1h02min.')
q27_e = doc.add_paragraph('(E) Em 2h15min.')

doc.add_paragraph('')

# ============= QUESTÃO 28 =============
doc.add_heading('Questão 28', level=3)
q28 = doc.add_paragraph('Giovana recortou um coração de 0,2 m² de uma cartolina, cujas medidas são dadas na imagem abaixo.\n\n')
q28.add_run('Quanto sobrou da cartolina?\n')

q28_a = doc.add_paragraph('(A) Sobrou 1,1 m² de cartolina.', style='List Number')
q28_b = doc.add_paragraph('(B) Sobrou 1,3 m² de cartolina.')
q28_c = doc.add_paragraph('(C) Sobrou 1,25 m² de cartolina.')
q28_d = doc.add_paragraph('(D) Sobrou 0,2 m² de cartolina.')
q28_e = doc.add_paragraph('(E) Sobrou 0,3 m² de cartolina.')

# Salvar documento
doc.save('/home/user/iases/Prova_Simulado_Questoes_15-28.docx')
print("Documento criado com sucesso: Prova_Simulado_Questoes_15-28.docx")
