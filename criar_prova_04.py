#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Criar documento
doc = Document()

# Configurar título
title = doc.add_heading('Concurso Público IASES 2022', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('AGENTE SOCIOEDUCATIVO - MASCULINO', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adicionar informações da prova
info = doc.add_paragraph()
info.add_run('Conhecimentos Específicos - Questões 43 a 56\n').bold = True
info.add_run('Simulado de Preparação').italic = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('_' * 80)

# ============= CONHECIMENTOS ESPECÍFICOS =============
doc.add_heading('Conhecimentos Específicos', level=2)

# ============= QUESTÃO 43 =============
doc.add_heading('Questão 43', level=3)
q43 = doc.add_paragraph('Em relação a Lei nº 9.455/1997, que define os crimes de tortura e dá outras providências, atribua V para verdadeiro e F para falso nas afirmativas abaixo.\n\n')

q43.add_run('(__) O crime de tortura é inafiançável e insuscetível de graça ou anistia.\n\n')
q43.add_run('(__) O disposto na Lei nº 9.455, de 7 de abril de 1997, aplica-se, ainda, quando o crime não tenha sido cometido em território nacional, sendo a vítima brasileira ou encontrando-se o agente em local sob jurisdição brasileira.\n\n')
q43.add_run('(__) No casos em que a tortura resultar lesão corporal de natureza grave ou gravíssima, a pena de reclusão será de oito a vinte anos; se resultar em morte, a pena de reclusão será de dez a trinta anos.\n\n')
q43.add_run('Assinale a alternativa que apresenta a sequência CORRETA.\n')

q43_a = doc.add_paragraph('(A) F, V, F.', style='List Number')
q43_b = doc.add_paragraph('(B) V, F, V.')
q43_c = doc.add_paragraph('(C) F, V, V.')
q43_d = doc.add_paragraph('(D) V, V, V.')
q43_e = doc.add_paragraph('(E) V, V, F.')

doc.add_paragraph('')

# ============= QUESTÃO 44 =============
doc.add_heading('Questão 44', level=3)
q44 = doc.add_paragraph('Em relação ao regime disciplinar dos deveres do Servidor Público, previstos na Lei Complementar Nº 46 de 31 de janeiro de 1994, assinale a alternativa que NÃO corresponde a um dever do Servidor Público.\n')

q44_a = doc.add_paragraph('(A) Observar as normas legais e regulamentares.', style='List Number')
q44_b = doc.add_paragraph('(B) Dar causa a sindicância ou processo administrativo-disciplinar, imputando a qualquer servidor público infração de que o sabe inocente.')
q44_c = doc.add_paragraph('(C) Guardar sigilo sobre assuntos da repartição.')
q44_d = doc.add_paragraph('(D) Representar contra ilegalidade, omissão ou abuso de poder, de que tenha tomado conhecimento, indicando elementos de prova para efeito de apuração em processo apropriado.')
q44_e = doc.add_paragraph('(E) Levar ao conhecimento da autoridade as irregularidades de que tiver ciência em razão do cargo ou função.')

doc.add_paragraph('')

# ============= QUESTÃO 45 =============
doc.add_heading('Questão 45', level=3)
q45 = doc.add_paragraph('Sobre o Sistema Único de Segurança Pública (SUSP), analise as afirmativas abaixo.\n\n')

q45.add_run('I. A segurança pública é atribuição de estados e municípios. Sendo a União responsável pela criação de diretrizes que serão compartilhadas em todo o país.\n\n')
q45.add_run('II. O SUSP cria uma arquitetura uniforme para a segurança pública em âmbito nacional, a partir de ações de compartilhamento de dados, operações integradas e colaborações nas estruturas de segurança pública federal, estadual e municipal.\n\n')
q45.add_run('III. O SUSP foi instituído pela Lei Nº 13.675 de 2018.\n\n')
q45.add_run('Fonte: gov.br/mj\n\n').italic = True
q45.add_run('É CORRETO o que se afirma em:\n')

q45_a = doc.add_paragraph('(A) I e III, apenas.', style='List Number')
q45_b = doc.add_paragraph('(B) II e III, apenas.')
q45_c = doc.add_paragraph('(C) I, II e III.')
q45_d = doc.add_paragraph('(D) I e II, apenas.')
q45_e = doc.add_paragraph('(E) III, apenas.')

doc.add_paragraph('')

# ============= QUESTÃO 46 =============
doc.add_heading('Questão 46', level=3)
q46 = doc.add_paragraph('No que se refere a Constituição Federal e o direito da Segurança Pública, analise as afirmativas abaixo.\n\n')

q46.add_run('I. Às Polícias Civis, dirigidas por delegados de polícia de carreira, incumbem, ressalvada a competência da União, as funções de Polícia Judiciária e a apuração de infrações penais, exceto as militares.\n\n')
q46.add_run('II. Às Polícias Militares cabem a polícia ostensiva e a preservação da ordem pública; aos corpos de bombeiros militares, além das atribuições definidas em lei, incumbe a execução de atividades de defesa civil.\n\n')
q46.add_run('III. À Polícia Rodoviária Federal, órgão permanente, organizado e mantido pelo Estado e estruturado em carreira, é destinado o patrulhamento sigiloso das rodovias estaduais.\n\n')
q46.add_run('É CORRETO o que se afirma em:\n')

q46_a = doc.add_paragraph('(A) II e III, apenas.', style='List Number')
q46_b = doc.add_paragraph('(B) I e III, apenas.')
q46_c = doc.add_paragraph('(C) I, II e III.')
q46_d = doc.add_paragraph('(D) II, apenas.')
q46_e = doc.add_paragraph('(E) I e II, apenas.')

doc.add_paragraph('')

# ============= QUESTÃO 47 =============
doc.add_heading('Questão 47', level=3)
q47 = doc.add_paragraph('No que diz respeito às revistas aos reclusos e inspeção de celas previstas nas Regras Mínimas das Nações Unidas para o Tratamento de Reclusos, é CORRETO afirmar que:\n')

q47_a = doc.add_paragraph('(A) Para fins de responsabilização, não é necessário que a administração prisional mantenha registos apropriados das revistas feitas aos reclusos e inspeções, em particular as que envolvem o ato de despir e de inspecionar partes íntimas do corpo e inspeções nas celas.', style='List Number')
q47_b = doc.add_paragraph('(B) As revistas íntimas invasivas devem ser conduzidas de forma privada e por pessoal treinado do sexo oposto que o recluso inspecionado.')
q47_c = doc.add_paragraph('(C) Os reclusos devem ter acesso aos documentos relacionados com os seus processos judiciais e ser autorizados a mantê-los consigo, desde que, a administração prisional tenha acesso a estes.')
q47_d = doc.add_paragraph('(D) As revistas aos reclusos e as inspeções podem ser utilizadas para intimidar ou invadir, desnecessariamente, a privacidade do recluso.')
q47_e = doc.add_paragraph('(E) As leis e regulamentos sobre as revistas aos reclusos e inspeções de celas devem estar em conformidade com as obrigações do Direito Internacional e devem ter em conta os padrões e as normas internacionais, uma vez considerada a necessidade de garantir a segurança dos estabelecimentos prisionais.')

doc.add_paragraph('')

# ============= QUESTÃO 48 =============
doc.add_heading('Questão 48', level=3)
q48 = doc.add_paragraph('No que se refere a atenção integral à saúde de adolescente em cumprimento de Medida Socioeducativa, prevista na Lei do Sistema Nacional de Atendimento Socioeducativo - SINASE, analise as afirmativas abaixo.\n\n')

q48.add_run('I. As entidades que oferecem programas de atendimento socioeducativo em meio aberto e de semiliberdade devem prestar orientações aos socioeducandos sobre o acesso aos serviços e às unidades do Sistema Único de Saúde (SUS).\n\n')
q48.add_run('II. As entidades que oferece programas de privação de liberdade devem contar com uma equipe mínima de profissionais de saúde cuja composição esteja em conformidade com as normas de referência do SUS.\n\n')
q48.add_run('III. É considerada uma diretriz da atenção integral à saúde do adolescente no Sistema de Atendimento Socioeducativo, a disponibilização de ações de atenção à saúde sexual e reprodutiva e à prevenção de doenças sexualmente transmissíveis.\n\n')
q48.add_run('É CORRETO o que se afirma em:\n')

q48_a = doc.add_paragraph('(A) III, apenas.', style='List Number')
q48_b = doc.add_paragraph('(B) I e II, apenas.')
q48_c = doc.add_paragraph('(C) I, II e III.')
q48_d = doc.add_paragraph('(D) I e III, apenas.')
q48_e = doc.add_paragraph('(E) II e III, apenas.')

doc.add_paragraph('')

# ============= QUESTÃO 49 =============
doc.add_heading('Questão 49', level=3)
q49 = doc.add_paragraph('Segundo o Estatuto da Criança e do Adolescente, é CORRETO afirmar que:\n')

q49_a = doc.add_paragraph('(A) Considera-se criança, a pessoa até onze anos de idade incompletos, e adolescente aquela entre doze e dezoito anos de idade.', style='List Number')
q49_b = doc.add_paragraph('(B) Considera-se criança, a pessoa até onze anos de idade incompletos, e adolescente aquela entre treze e dezoito anos de idade.')
q49_c = doc.add_paragraph('(C) Considera-se criança, a pessoa até onze anos de idade incompletos, e adolescente aquela entre doze e vinte e um anos de idade.')
q49_d = doc.add_paragraph('(D) Considera-se criança, a pessoa até dez anos de idade incompletos, e adolescente aquela entre treze e dezoito anos de idade.')
q49_e = doc.add_paragraph('(E) Considera-se criança, a pessoa até doze anos de idade incompletos, e adolescente aquela entre doze e dezoito anos de idade.')

doc.add_paragraph('')

# ============= QUESTÃO 50 =============
doc.add_heading('Questão 50', level=3)
obs50 = doc.add_paragraph()
obs50.add_run('(Questão anulada)\n\n').italic = True
q50 = doc.add_paragraph('Em relação a composição do Conselho Nacional dos Direitos da Criança e do Adolescente - CONANDA, assinale a alternativa CORRETA.\n\n')
q50.add_run('Fonte: gov.br/mdh\n\n').italic = True

q50_a = doc.add_paragraph('(A) O CONANDA é órgão colegiado de composição distinta, integrado por quatorze representantes do Poder Executivo, assegurada a participação dos órgãos executores das políticas sociais básicas e, em igual número, por representantes de entidades não-governamentais de âmbito nacional de promoção, proteção, defesa e controle social da política de atendimento dos direitos da criança e do adolescente.', style='List Number')
q50_b = doc.add_paragraph('(B) O CONANDA é órgão colegiado de composição paritária, integrado por vinte e oito representantes do Poder Executivo, assegurada a participação dos órgãos executores das políticas sociais básicas e, em igual número, por representantes de entidades não-governamentais de âmbito nacional de promoção, proteção, defesa e controle social da política de atendimento dos direitos da criança e do adolescente.')
q50_c = doc.add_paragraph('(C) O CONANDA é órgão colegiado de composição distinta, integrado por vinte e oito representantes do Poder Executivo, assegurada a participação dos órgãos executores das políticas sociais básicas e, em igual número, por representantes de entidades não-governamentais de âmbito nacional de promoção, proteção, defesa e controle social da política de atendimento dos direitos da criança e do adolescente.')
q50_d = doc.add_paragraph('(D) O CONANDA é órgão colegiado de composição distinta, integrado por vinte e oito representantes do Poder Executivo, assegurada a participação dos órgãos executores das políticas sociais básicas e, por quatorze representantes de entidades não-governamentais de âmbito nacional de promoção, proteção, defesa e controle social da política de atendimento dos direitos da criança e do adolescente.')
q50_e = doc.add_paragraph('(E) O CONANDA é órgão colegiado de composição paritária, integrado por quatorze representantes do Poder Executivo, assegurada a participação dos órgãos executores das políticas sociais básicas e, em igual número, por representantes de entidades não-governamentais de âmbito nacional de promoção, proteção, defesa e controle social da política de atendimento dos direitos da criança e do adolescente.')

doc.add_paragraph('')

# ============= QUESTÃO 51 =============
doc.add_heading('Questão 51', level=3)
q51 = doc.add_paragraph('Segundo a Lei Nº 9.455, de 7 de abril de 1997, a pena para o crime de tortura que consiste em submeter alguém, sob sua guarda, poder ou autoridade, com emprego de violência ou grave ameaça, a intenso sofrimento físico ou mental, como forma de aplicar castigo pessoal ou medida de caráter preventivo, é de:\n')

q51_a = doc.add_paragraph('(A) Reclusão, de um a cinco anos.', style='List Number')
q51_b = doc.add_paragraph('(B) Reclusão, de quatro a doze anos.')
q51_c = doc.add_paragraph('(C) Reclusão, de um a doze anos.')
q51_d = doc.add_paragraph('(D) Reclusão, de três a dez anos.')
q51_e = doc.add_paragraph('(E) Reclusão, de dois a oito anos.')

doc.add_paragraph('')

# ============= QUESTÃO 52 =============
doc.add_heading('Questão 52', level=3)
q52 = doc.add_paragraph('São direitos dos adolescentes submetidos ao cumprimento de medida socioeducativa, sem prejuízo de outros, previstos na Lei do Sistema Nacional de Atendimento Socioeducativo - SINASE, EXCETO:\n')

q52_a = doc.add_paragraph('(A) Ser informado, inclusive por escrito, das normas de organização e funcionamento do programa de atendimento e também das previsões de natureza disciplinar.', style='List Number')
q52_b = doc.add_paragraph('(B) Peticionar, por escrito ou verbalmente, diretamente a qualquer autoridade ou órgão público, devendo, obrigatoriamente, ser respondido em até 15 dias.')
q52_c = doc.add_paragraph('(C) Ter atendimento garantido em creche e pré-escola aos filhos de 0 a 5 anos.')
q52_d = doc.add_paragraph('(D) Ser incluído em programa de meio aberto quando inexistir vaga para o cumprimento de medida de privação da liberdade, inclusive nos casos de ato infracional cometido mediante grave ameaça ou violência à pessoa, quando o adolescente deverá ser internado em Unidade mais próxima de seu local de residência.')
q52_e = doc.add_paragraph('(E) Receber, sempre que solicitar, informações sobre a evolução de seu plano individual, participando, obrigatoriamente, de sua elaboração e, se for o caso, reavaliação.')

doc.add_paragraph('')

# ============= QUESTÃO 53 =============
doc.add_heading('Questão 53', level=3)
q53 = doc.add_paragraph('Em relação a Composição do Sistema Único de Segurança Pública (SUSP), assinale a alternativa INCORRETA.\n')

q53_a = doc.add_paragraph('(A) São integrantes estratégicos do SUSP: a União, os Estados, o Distrito Federal e os Municípios, por intermédio dos respectivos Poderes Executivos; os Conselhos de Segurança Pública e Defesa Social dos três entes federados.', style='List Number')
q53_b = doc.add_paragraph('(B) Não são considerados integrantes operacionais do SUSP, a Polícia Federal e a Polícia Rodoviária Federal.')
q53_c = doc.add_paragraph('(C) Os sistemas estaduais, distrital e municipais são responsáveis pela implementação dos respectivos programas, ações e projetos de segurança pública, com liberdade de organização e funcionamento.')
q53_d = doc.add_paragraph('(D) O Sistema Único de Segurança Pública é integrado pelos órgãos de que trata o art. 144 da Constituição Federal, pelos Agentes Penitenciários, pelas Guardas Municipais e pelos demais integrantes estratégicos e operacionais, que atuarão nos limites de suas competências, de forma cooperativa, sistêmica e harmônica.')
q53_e = doc.add_paragraph('(E) O Sistema Único de Segurança Pública, tem como órgão central o Ministério Extraordinário da Segurança Pública.')

doc.add_paragraph('')

# ============= QUESTÃO 54 =============
doc.add_heading('Questão 54', level=3)
q54 = doc.add_paragraph('Em relação a Constituição Federal e a administração pública, atribua V para verdadeiro e F para falso nas afirmativas abaixo.\n\n')

q54.add_run('(__) O prazo de validade do concurso público será de até dois anos, prorrogável uma vez, por igual período.\n\n')
q54.add_run('(__) Durante o prazo improrrogável previsto no edital de convocação, aquele aprovado em concurso público de provas ou de provas e títulos será convocado com prioridade sobre novos concursados para assumir cargo ou emprego, na carreira.\n\n')
q54.add_run('(__) A investidura em cargo ou emprego público depende de aprovação prévia em concurso público de provas ou de provas e títulos, não sendo permitido a nomeações para cargo em comissão, mesmo declarado em lei de livre nomeação e exoneração.\n\n')
q54.add_run('Assinale a alternativa que apresenta a sequência CORRETA.\n')

q54_a = doc.add_paragraph('(A) F, F, F.', style='List Number')
q54_b = doc.add_paragraph('(B) V, V, V.')
q54_c = doc.add_paragraph('(C) F, V, F.')
q54_d = doc.add_paragraph('(D) V, V, F.')
q54_e = doc.add_paragraph('(E) V, F, V.')

doc.add_paragraph('')

# ============= QUESTÃO 55 =============
doc.add_heading('Questão 55', level=3)
q55 = doc.add_paragraph('Referente ao Processo Administrativo-Disciplinar e, em relação as irregularidades no serviço público, previsto na Lei Complementar Nº 46 de 31 de janeiro de 1994, analise as afirmativas abaixo.\n\n')

q55.add_run('I. As denúncias sobre irregularidades serão objeto de apuração, desde que, contenham a identificação do denunciante, devendo ser formuladas por escrito.\n\n')
q55.add_run('II. A sindicância se constituirá de averiguação sumária promovida no intuito de obter informações ou esclarecimentos necessários à determinação do verdadeiro significado dos fatos denunciados.\n\n')
q55.add_run('III. A autoridade que tiver ciência de irregularidade no serviço público é obrigada a promover a sua apuração imediata, mediante sindicância ou processo administrativo-disciplinar, assegurada ao denunciado ampla defesa.\n\n')
q55.add_run('É CORRETO o que se afirma em:\n')

q55_a = doc.add_paragraph('(A) I e III, apenas.', style='List Number')
q55_b = doc.add_paragraph('(B) II e III, apenas.')
q55_c = doc.add_paragraph('(C) I e II, apenas.')
q55_d = doc.add_paragraph('(D) II, apenas.')
q55_e = doc.add_paragraph('(E) I, II e III.')

doc.add_paragraph('')

# ============= QUESTÃO 56 =============
doc.add_heading('Questão 56', level=3)
q56 = doc.add_paragraph('De acordo com a Lei do Sistema Nacional de Atendimento Socioeducativo - SINASE, compete à União, EXCETO:\n')

q56_a = doc.add_paragraph('(A) Financiar, com os demais entes federados, a execução de programas e serviços do Sinase.', style='List Number')
q56_b = doc.add_paragraph('(B) Contribuir para a qualificação e ação em rede dos Sistemas de Atendimento Socioeducativo.')
q56_c = doc.add_paragraph('(C) Editar normas complementares para a organização e funcionamento do seu sistema de atendimento e dos sistemas municipais.')
q56_d = doc.add_paragraph('(D) Elaborar o Plano Nacional de Atendimento Socioeducativo, em parceria com os Estados, o Distrito Federal e os Municípios.')
q56_e = doc.add_paragraph('(E) Formular e coordenar a execução da política nacional de atendimento socioeducativo.')

# Salvar documento
doc.save('/home/user/iases/Prova_Simulado_Questoes_43-56.docx')
print("Documento criado com sucesso: Prova_Simulado_Questoes_43-56.docx")
