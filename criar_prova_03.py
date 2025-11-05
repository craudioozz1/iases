#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Criar documento
doc = Document()

# Configurar título
title = doc.add_heading('Concurso Público IASES 2022', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('AGENTE SOCIOEDUCATIVO - MASCULINO', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adicionar informações da prova
info = doc.add_paragraph()
info.add_run('Matemática (Q29-30), Informática (Q31-40) e Conhecimentos Específicos (Q41-42)\n').bold = True
info.add_run('Questões 29 a 42 - Simulado de Preparação').italic = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('_' * 80)

# ============= MATEMÁTICA =============
doc.add_heading('Matemática', level=2)

# ============= QUESTÃO 29 =============
doc.add_heading('Questão 29', level=3)
q29 = doc.add_paragraph('A professora Ana vai fazer uma excursão com seus 30 alunos e, para organizar, vai formar duplas para distribuir os lugares no ônibus. De quantas maneiras diferentes ela pode formar essas duplas?\n')

q29_a = doc.add_paragraph('(A) De 560 maneiras diferentes.', style='List Number')
q29_b = doc.add_paragraph('(B) De 920 maneiras diferentes.')
q29_c = doc.add_paragraph('(C) De 630 maneiras diferentes.')
q29_d = doc.add_paragraph('(D) De 755 maneiras diferentes.')
q29_e = doc.add_paragraph('(E) De 870 maneiras diferentes.')

doc.add_paragraph('')

# ============= QUESTÃO 30 =============
doc.add_heading('Questão 30', level=3)
q30 = doc.add_paragraph('A boia sinalizadora é um artefato utilizado em praias, represas, enseadas, marinas, rios e lagos que necessitem de sinalização e para ser visualizada facilmente à noite possui uma lâmpada pisca-pisca. Na entrada de uma enseada foram colocadas duas boias sinalizadoras separadas por 50 metros e que piscam em frequências diferentes. Uma boia pisca 10 vezes por minuto e a outra boia pisca 15 vezes por minuto. Sabendo-se que num certo momento as luzes das boias piscam ao mesmo tempo, após quantos segundos elas voltarão a piscar simultaneamente?\n')

q30_a = doc.add_paragraph('(A) 15 segundos.', style='List Number')
q30_b = doc.add_paragraph('(B) 25 segundos.')
q30_c = doc.add_paragraph('(C) 20 segundos.')
q30_d = doc.add_paragraph('(D) 12 segundos.')
q30_e = doc.add_paragraph('(E) 18 segundos.')

doc.add_page_break()

# ============= INFORMÁTICA =============
doc.add_heading('Informática', level=2)

# ============= QUESTÃO 31 =============
doc.add_heading('Questão 31', level=3)
q31 = doc.add_paragraph('Um sistema operacional é um software que se posiciona entre a pessoa usuária e os componentes físicos de um computador (Hardwares). Por meio do Sistema Operacional é possível controlar a execução de tarefas e programas, assim como o gerenciamento da memória, dispositivos e arquivos. Sobre Sistema Operacional, assinale a alternativa CORRETA:\n')

q31_a = doc.add_paragraph('(A) O Kernel é o componente central de um sistema operacional. Ele opera no núcleo do computador, garantindo que haja comunicação entre os componentes do Hardware e o terminal no qual o sistema operacional é executado.', style='List Number')
q31_b = doc.add_paragraph('(B) O sistema operacional cria uma comunicação direta por meio de clusters com os dispositivos conectados no Hardware. Um cluster, portanto, é um software que viabiliza a execução de um dispositivo, que pode ser um pen drive, mouse, teclado ou até mesmo a placa de vídeo da máquina.')
q31_c = doc.add_paragraph('(C) A interface do usuário é o software que interage com o sistema operacional e os aplicativos executados em um computador. O sistema operacional e os aplicativos necessitam desta interface para serem executados.')
q31_d = doc.add_paragraph('(D) O Firmware é o nome dado a um Sistema Operacional que contém um conjunto de instruções programadas em um dispositivo Hardware, responsável por providenciar que o funcionamento do dispositivo seja devidamente comunicado e executado em outro componente.')
q31_e = doc.add_paragraph('(E) Os Sistemas operacionais não são responsáveis por manter registros da performance do sistema, este serviço é gerenciado pelos softwares utilitários caso ocorra algum imprevisto ou demora na execução de programas.')

doc.add_paragraph('')

# ============= QUESTÃO 32 =============
doc.add_heading('Questão 32', level=3)
q32 = doc.add_paragraph('Saber numerar páginas no editor de texto Word do Microsoft Office é algo essencial para quem está escrevendo trabalhos acadêmicos e documentos que exigem a numeração. Sobre inserção da numeração de página, analise as afirmativas a seguir:\n\n')

q32.add_run('I. Para definir onde será o começo da numeração de forma automática no Word sem precisar editar todas as páginas, basta ir na aba "Inserir", clicar em "Número de Página" para ver mais opções; em seguida, clicar em "Formatar Número de Página"; feito isso, ir na categoria "Numeração da página" marcar a opção "Iniciar em" para definir onde será o começo da numeração.\n\n')
q32.add_run('II. Entre suas opções de formatação, o editor Word do Microsoft Office oferece várias ferramentas para a enumeração de páginas de maneira rápida e automática. Para a numeração personalizada é necessário instalar o suplemento Word Page Plus.\n\n')
q32.add_run('III. Para numerar páginas no Microsoft Word, basta ir na barra de menu principal, clicar em "Inserir"; em seguida, clicar em "Número de Página" e selecionar um dos layouts disponíveis para numeração. Seguindo estes passos as páginas serão numeradas automaticamente.\n\n')
q32.add_run('É CORRETO afirmar que:\n')

q32_a = doc.add_paragraph('(A) Apenas a afirmativa II é verdadeira.', style='List Number')
q32_b = doc.add_paragraph('(B) Apenas as afirmativas I e III são verdadeiras.')
q32_c = doc.add_paragraph('(C) Apenas as afirmativas I e II são verdadeiras.')
q32_d = doc.add_paragraph('(D) Apenas a afirmativa I é verdadeira.')
q32_e = doc.add_paragraph('(E) As afirmativas I, II e III são verdadeiras.')

doc.add_paragraph('')

# ============= QUESTÃO 33 =============
doc.add_heading('Questão 33', level=3)
obs33 = doc.add_paragraph()
obs33.add_run('(Questão anulada)\n\n').italic = True
q33 = doc.add_paragraph('Assinale a alternativa que corresponda ao atalho utilizado, no sistema operacional Windows, para selecionar um bloco de texto.\n')

q33_a = doc.add_paragraph('(A) Tecla do logo tipo Windows.', style='List Number')
q33_b = doc.add_paragraph('(B) Alt + Barra de espaço.')
q33_c = doc.add_paragraph('(C) Ctrl + Shift + Qualquer tecla de seta.')
q33_d = doc.add_paragraph('(D) Ctrl + Esc.')
q33_e = doc.add_paragraph('(E) Tab + Ctrl + B.')

doc.add_paragraph('')

# ============= QUESTÃO 34 =============
doc.add_heading('Questão 34', level=3)
q34 = doc.add_paragraph('No Microsoft Excel a formatação condicional facilita realçar certos valores ou tornar determinadas células fáceis de identificar. Isso altera a aparência de um intervalo de células com base em uma condição (ou critérios). Analise as afirmativas a seguir sobre formatação condicional:\n\n')

q34.add_run('I. Para criar uma regra de formatação condicional personalizada o usuário do Excel deverá selecionar o intervalo de células, a tabela ou a planilha inteira; na guia Página Inicial, clicar em Formatação Condicional e depois em Nova Regra; Escolher um estilo, por exemplo, escala de 3 cores, selecionar as condições que deseja e, em seguida, clicar em OK.\n\n')
q34.add_run('II. Para realçar valores em células específicas (como exemplos datas depois de uma determinada semana, números entre 50 e 100 ou os 10% inferiores dos resultados), o usuário deverá apontar para Realçar Regras de Células ou Regras de Primeiros/Últimos e, em seguida, clicar na opção adequada.\n\n')
q34.add_run('III. Para realizar uso de formatações condicionais o usuário pode inserir a função =FCOND() na célula. Por exemplo: =FCOND(LIN();2)=0, o resultado seria VERDADEIRO ou FALSO, par ou ímpar, pois a função FCOND retorna o resto da divisão do número da célula por 2 e verifica se o retorno é 0.\n\n')
q34.add_run('Assinale a alternativa correta:\n')

q34_a = doc.add_paragraph('(A) Apenas as afirmativas II e III são verdadeiras.', style='List Number')
q34_b = doc.add_paragraph('(B) Apenas a afirmativa II é verdadeira.')
q34_c = doc.add_paragraph('(C) Apenas a afirmativa III é verdadeira.')
q34_d = doc.add_paragraph('(D) Apenas as afirmativas I e II são verdadeiras.')
q34_e = doc.add_paragraph('(E) Apenas as afirmativas I e III são verdadeiras.')

doc.add_paragraph('')

# ============= QUESTÃO 35 =============
doc.add_heading('Questão 35', level=3)
obs35 = doc.add_paragraph()
obs35.add_run('(Questão anulada)\n\n').italic = True
q35 = doc.add_paragraph('O E-Docs é o Sistema corporativo de gestão de documentos que possibilita uma interação entre estado e sociedade de maneira mais objetiva e acessível. Analise as afirmativas abaixo, no que se refere às funcionalidades do E-Docs:\n\n')

q35.add_run('I. Para acessar, especificamente, os documentos que estão solicitando assinatura, o usuário deverá selecionar a opção "Para Assinar", clicar no link do documento que será assinado, depois clicar em "assinar" e por fim clicar novamente em "assinar" para confirmar a assinatura.\n\n')
q35.add_run('II. Por meio do E-Docs, todos os cidadãos e cidadãs que realizarem o cadastro no sistema, podem encaminhar documentos, enviar relatórios, informar o andamento dos projetos e assinar documentos. Para o acompanhamento de editais, parcerias e contratos, o usuário deve no ato da habilitação do login, solicitar uma senha super usuário.\n\n')
q35.add_run('III. Para localizar um documento enviado para o estado, o usuário deve acessar a conta no E-Docs, clicar em "Ir para encaminhamentos" e, todos os documentos recebidos poderão ser encontrados na "Caixa de Entrada" e todos os documentos enviados poderão ser encontrados na caixa de saída.\n\n')
q35.add_run('É CORRETO afirmar que:\n')

q35_a = doc.add_paragraph('(A) As afirmativas I, II e III são verdadeiras.', style='List Number')
q35_b = doc.add_paragraph('(B) Apenas a afirmativa II é verdadeira.')
q35_c = doc.add_paragraph('(C) Apenas as afirmativas I e II são verdadeiras.')
q35_d = doc.add_paragraph('(D) Apenas as afirmativas I e III são verdadeiras.')
q35_e = doc.add_paragraph('(E) Apenas a afirmativa I é verdadeira.')

doc.add_paragraph('')

# ============= QUESTÃO 36 =============
doc.add_heading('Questão 36', level=3)
obs36 = doc.add_paragraph()
obs36.add_run('(Questão anulada)\n\n').italic = True
q36 = doc.add_paragraph('Em um mundo conectado pela internet, a cibersegurança é um cuidado cada vez mais importante para usuário e empresas. Para se adaptar a esse cenário de conectividade e ameaças virtuais, torna-se indispensável o conhecimento básico sobre antivírus.\n\n')
q36.add_run('Sobre os antivírus, analise as afirmativas a seguir:\n\n')

q36.add_run('I. Alguns antivírus podem ser específicos para um determinado tipo de vírus, como os antispywares, que são focados em combater spywares e adwares - dois tipos de vírus que contaminam arquivos executáveis, arquivos encriptados e kernel do sistema operacional.\n\n')
q36.add_run('II. Antivírus é um software que identifica e protege os dispositivos de malwares, também conhecidos como vírus. Esse programa pode ser instalado em computadores e dispositivos móveis, como celulares e tablets.\n\n')
q36.add_run('III. No antivírus, a quarentena é um espaço de proteção criptografado e gerenciado pelo antivírus, para que o possível vírus não se espalhe pelo sistema operacional do dispositivo. Arquivos e programas são encaminhados para a quarentena quando o antivírus ainda não identificou exatamente o tipo de vírus ou problema apresentado.\n\n')
q36.add_run('Assinale a alternativa correta:\n')

q36_a = doc.add_paragraph('(A) Apenas a afirmativa I é verdadeira.', style='List Number')
q36_b = doc.add_paragraph('(B) Apenas as afirmativas II e III são verdadeiras.')
q36_c = doc.add_paragraph('(C) Apenas a afirmativa III é verdadeira.')
q36_d = doc.add_paragraph('(D) Apenas as afirmativas I e II são verdadeiras.')
q36_e = doc.add_paragraph('(E) As afirmativas I, II e III são verdadeiras.')

doc.add_paragraph('')

# ============= QUESTÃO 37 =============
doc.add_heading('Questão 37', level=3)
q37 = doc.add_paragraph('"Um firewall é um sistema de segurança de rede de computadores que restringe o tráfego da Internet para, de ou em uma rede privada".\n\n')
q37.add_run('(https://www.kaspersky.com.br/resource-center/definitions/firewall)\n\n').italic = True
q37.add_run('Sobre o funcionamento e características do firewall, analise as seguintes afirmativas:\n\n')

q37.add_run('I. O firewall decide qual tráfego de rede pode passar e qual tráfego é considerado perigoso. Basicamente, ele atua como um filtro, separando o que é bom do que é ruim, o confiável do não confiável.\n\n')
q37.add_run('II. Firewalls de host ou "firewalls de software" envolvem a aplicação de um ou mais firewalls entre redes externas e redes privadas internas. Eles regulam o tráfego de rede de entrada e saída, separando redes públicas externas, como a Internet global, de redes internas, como redes Wi-Fi domésticas, intranets corporativas ou intranets nacionais.\n\n')
q37.add_run('III. Os firewalls de proxy, também conhecidos como firewalls de nível de aplicação, são únicos no que se refere à leitura e à filtragem de protocolos de aplicativos. Eles combinam inspeção em nível de aplicação, ou "inspeção profunda de pacotes (DPI)" e inspeção com estado.\n\n')
q37.add_run('É CORRETO afirmar que:\n')

q37_a = doc.add_paragraph('(A) Apenas a afirmativa III é verdadeira.', style='List Number')
q37_b = doc.add_paragraph('(B) Apenas as afirmativas I e III são verdadeiras.')
q37_c = doc.add_paragraph('(C) Apenas as afirmativas I e II são verdadeiras.')
q37_d = doc.add_paragraph('(D) As afirmativas I, II e III são verdadeiras.')
q37_e = doc.add_paragraph('(E) Apenas a afirmativa I é verdadeira.')

doc.add_paragraph('')

# ============= QUESTÃO 38 =============
doc.add_heading('Questão 38', level=3)
q38 = doc.add_paragraph('Um site de busca ou buscador é, em termos gerais, um sistema encarregado de pesquisar arquivos armazenados em servidores da Internet. Nesse sentido, analise as afirmativas a seguir:\n\n')

q38.add_run('I. Para encontrar o resultado de pesquisas, os buscadores recorrem à identificação da palavra-chave usada pelo usuário durante sua pesquisa e, como resultado, entregam uma lista de links que direcionam a sites que mencionam assuntos relacionados ao termo pesquisado.\n\n')
q38.add_run('II. Os sites de busca são classificados, principalmente, em buscadores hierárquicos; diretórios e metabuscadores.\n\n')
q38.add_run('III. Os Metabuscadores são interfaces que buscam as pesquisas através de um banco de dados, analisam e encaminham as pesquisa para servidores.\n\n')
q38.add_run('IV. O AsK é o sistema de pesquisa nativo que dispositivos da Microsoft utilizam desde o lançamento do Windows 8.\n\n')
q38.add_run('É CORRETO afirmar que:\n')

q38_a = doc.add_paragraph('(A) Apenas as afirmativas I e IV são verdadeiras.', style='List Number')
q38_b = doc.add_paragraph('(B) Apenas as afirmativas I e II são verdadeiras.')
q38_c = doc.add_paragraph('(C) Apenas as afirmativas II e IV são verdadeiras.')
q38_d = doc.add_paragraph('(D) As afirmativas I, II, III e IV são verdadeiras.')
q38_e = doc.add_paragraph('(E) Apenas a afirmativa III é verdadeira.')

doc.add_paragraph('')

# ============= QUESTÃO 39 =============
doc.add_heading('Questão 39', level=3)
q39 = doc.add_paragraph('São considerados programas antivírus, EXCETO:\n')

q39_a = doc.add_paragraph('(A) MySQL.', style='List Number')
q39_b = doc.add_paragraph('(B) AVG.')
q39_c = doc.add_paragraph('(C) Kaspersky.')
q39_d = doc.add_paragraph('(D) Avast.')
q39_e = doc.add_paragraph('(E) McAfee.')

doc.add_paragraph('')

# ============= QUESTÃO 40 =============
doc.add_heading('Questão 40', level=3)
q40 = doc.add_paragraph('Sobre a URL, é CORRETO afirmar que:\n')

q40_a = doc.add_paragraph('(A) Consiste em uma ferramenta de rede que atua como intermediário entre dispositivos cliente e servidores de destino.', style='List Number')
q40_b = doc.add_paragraph('(B) Consiste em um protocolo de transferência que possibilita que as pessoas que inserem a URL do seu site na Web possam ver os conteúdos e dados que nele existem.')
q40_c = doc.add_paragraph('(C) Consiste em uma unidade identificadora que fornece uma maneira lógica de acessar informações presentes em computadores remotos, como um servidor da web ou um site de armazenamento em nuvem.')
q40_d = doc.add_paragraph('(D) Consiste em um elemento de hipermídia formado por um trecho de texto em destaque ou por um elemento gráfico que, ao ser acionado, provoca a exibição de novo hiperdocumento.')
q40_e = doc.add_paragraph('(E) Consiste em um local na Internet identificado por um nome de domínio, constituído por uma ou mais páginas de hipertexto, que podem conter textos, gráficos e informações em multimídia.')

doc.add_page_break()

# ============= CONHECIMENTOS ESPECÍFICOS =============
doc.add_heading('Conhecimentos Específicos', level=2)

# ============= QUESTÃO 41 =============
doc.add_heading('Questão 41', level=3)
q41 = doc.add_paragraph('Em relação às informações e direito de reclamação dos reclusos, descritas nas Regras Mínimas das Nações Unidas para o Tratamento de Reclusos, assinale a alternativa INCORRETA.\n')

q41_a = doc.add_paragraph('(A) Todo recluso deve ter a oportunidade de, em qualquer dia, formular pedidos ou reclamações ao diretor do estabelecimento prisional ou ao membro do pessoal prisional autorizado a representá-lo.', style='List Number')
q41_b = doc.add_paragraph('(B) Todo recluso, no momento da admissão, deve receber informação escrita sobre a legislação e os regulamentos do estabelecimento prisional e do sistema prisional.')
q41_c = doc.add_paragraph('(C) Todo recluso deve ter o direito de fazer um pedido ou reclamação sobre seu tratamento, com censura quanto ao conteúdo, à administração prisional central, à autoridade judicial ou a outras autoridades competentes, incluindo os que têm poderes de revisão e de reparação.')
q41_d = doc.add_paragraph('(D) Todo pedido ou reclamação deve ser prontamente apreciado e respondido sem demora. Se o pedido ou a reclamação for rejeitado, ou no caso de atraso indevido, o reclamante deve ter o direito de apresentá-lo à autoridade judicial ou a outra autoridade.')
q41_e = doc.add_paragraph('(E) Se recluso for analfabeto, as informações devem ser-lhe comunicadas oralmente. Os reclusos com deficiências sensoriais devem receber as informações de forma apropriada às suas necessidades.')

doc.add_paragraph('')

# ============= QUESTÃO 42 =============
doc.add_heading('Questão 42', level=3)
q42 = doc.add_paragraph('Segundo o Ministério dos Direitos Humanos e da Cidadania, são atribuições do Conselho Nacional dos Direitos da Criança e do Adolescente - CONANDA, EXCETO:\n\n')
q42.add_run('Fonte: gov.br/mdh\n\n').italic = True

q42_a = doc.add_paragraph('(A) Acompanhar a elaboração e a execução do orçamento da União, verificando se estão assegurados os recursos necessários para a execução das políticas de promoção e defesa dos direitos da população infanto-juvenil.', style='List Number')
q42_b = doc.add_paragraph('(B) Definir as diretrizes para a criação e o funcionamento dos Conselhos Estaduais, Distrital e Municipais dos Direitos da Criança e do Adolescente e dos Conselhos Tutelares.')
q42_c = doc.add_paragraph('(C) Estimular, apoiar e promover a manutenção de bancos de dados com informações sobre a infância e a adolescência.')
q42_d = doc.add_paragraph('(D) Fiscalizar as ações de promoção dos direitos da infância e adolescência executadas por organismos governamentais e não-governamentais.')
q42_e = doc.add_paragraph('(E) Convocar, a cada cinco anos, a Conferência Nacional de Saúde e dos Direitos da Criança e do Adolescente.')

# Salvar documento
doc.save('/home/user/iases/Prova_Simulado_Questoes_29-42.docx')
print("Documento criado com sucesso: Prova_Simulado_Questoes_29-42.docx")
